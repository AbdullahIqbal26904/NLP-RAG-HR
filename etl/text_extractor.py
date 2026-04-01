"""
Text Extraction Layer for Resume Parser ETL System.

This module handles raw text extraction from PDF and DOCX files.
It preserves original formatting and does NOT perform aggressive cleaning.
Cleaning is delegated to the Text Cleaning Layer.

Supported formats:
    - PDF (text-based, via pdfplumber)
    - DOCX (via python-docx)
"""

from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Optional
import logging

import pdfplumber
from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from .layout_analyzer import LayoutAnalyzer, extract_page_with_layout, LayoutInfo

logger = logging.getLogger(__name__)


class FileType(Enum):
    """Supported file types for text extraction."""
    PDF = "pdf"
    DOCX = "docx"


class ExtractionStatus(Enum):
    """Status of the extraction operation."""
    SUCCESS = "success"
    PARTIAL = "partial"  # Some pages failed but others succeeded
    FAILED = "failed"
    EMPTY = "empty"  # File has no extractable text


@dataclass
class ExtractionResult:
    """
    Result of text extraction operation.
    
    Attributes:
        text: The extracted raw text (None if extraction failed completely)
        status: Status of the extraction operation
        file_path: Original file path
        file_type: Detected file type
        page_count: Total number of pages (for PDFs)
        pages_extracted: Number of pages successfully extracted
        error_message: Error details if extraction failed
        warnings: List of non-fatal issues encountered
    """
    text: Optional[str]
    status: ExtractionStatus
    file_path: str
    file_type: FileType
    page_count: int = 0
    pages_extracted: int = 0
    error_message: Optional[str] = None
    warnings: list[str] = field(default_factory=list)

    @property
    def is_success(self) -> bool:
        """Check if extraction was successful (full or partial)."""
        return self.status in (ExtractionStatus.SUCCESS, ExtractionStatus.PARTIAL)

    @property
    def has_text(self) -> bool:
        """Check if any text was extracted."""
        return self.text is not None and len(self.text.strip()) > 0


class TextExtractorError(Exception):
    """Base exception for text extraction errors."""
    pass


class FileNotFoundError(TextExtractorError):
    """Raised when the input file does not exist."""
    pass


class UnsupportedFileTypeError(TextExtractorError):
    """Raised when the file type is not supported."""
    pass


class CorruptedFileError(TextExtractorError):
    """Raised when the file is corrupted or cannot be read."""
    pass


class BaseTextExtractor(ABC):
    """
    Abstract base class for text extractors.
    
    All format-specific extractors must implement this interface
    to ensure consistent behavior across the extraction layer.
    """

    @abstractmethod
    def extract(self) -> ExtractionResult:
        """
        Extract text from the file.
        
        Returns:
            ExtractionResult containing the extracted text and metadata.
        """
        pass

    @abstractmethod
    def can_extract(self) -> bool:
        """
        Check if the file can be extracted (exists, is readable, etc.).
        
        Returns:
            True if extraction is possible, False otherwise.
        """
        pass


class PdfTextExtractor(BaseTextExtractor):
    """
    Layout-aware text extractor for PDF files using pdfplumber.
    
    Handles both single-column and multi-column PDF layouts.
    Multi-column resumes (sidebars, two-column templates) are detected
    automatically and each column is extracted independently to preserve
    semantic structure.
    
    For scanned PDFs (image-based), this extractor will return empty text
    with appropriate warnings. OCR integration can be added as a future
    enhancement.
    
    Attributes:
        file_path: Path to the PDF file.
    """

    # Page separator to maintain structure when combining pages
    PAGE_SEPARATOR = "\n\n--- PAGE BREAK ---\n\n"

    def __init__(self, file_path: str):
        """
        Initialize the PDF extractor.
        
        Args:
            file_path: Path to the PDF file to extract.
        """
        self._file_path = Path(file_path)
        self._layout_analyzer = LayoutAnalyzer()

    def can_extract(self) -> bool:
        """Check if the PDF file exists and is readable."""
        return self._file_path.exists() and self._file_path.is_file()

    def extract(self) -> ExtractionResult:
        """
        Extract text from all pages of the PDF.
        
        Returns:
            ExtractionResult with extracted text and metadata.
            
        Note:
            - Preserves original text formatting (no aggressive cleaning)
            - Pages are separated by PAGE_SEPARATOR for structure preservation
            - Empty pages are logged as warnings but don't fail extraction
        """
        if not self.can_extract():
            return ExtractionResult(
                text=None,
                status=ExtractionStatus.FAILED,
                file_path=str(self._file_path),
                file_type=FileType.PDF,
                error_message=f"File not found or not readable: {self._file_path}"
            )

        extracted_pages: list[str] = []
        warnings: list[str] = []
        page_count = 0
        pages_extracted = 0

        try:
            with pdfplumber.open(self._file_path) as pdf:
                page_count = len(pdf.pages)
                
                if page_count == 0:
                    return ExtractionResult(
                        text=None,
                        status=ExtractionStatus.EMPTY,
                        file_path=str(self._file_path),
                        file_type=FileType.PDF,
                        page_count=0,
                        error_message="PDF has no pages"
                    )

                for page_num, page in enumerate(pdf.pages, start=1):
                    try:
                        # Layout-aware extraction: detects multi-column
                        # layouts and extracts each column independently
                        page_text, layout = extract_page_with_layout(
                            page, self._layout_analyzer
                        )

                        if layout.is_multi_column:
                            logger.info(
                                "Page %d of %s: detected %d-column layout",
                                page_num, self._file_path, layout.column_count,
                            )
                            warnings.extend(
                                f"Page {page_num}: {w}" for w in layout.warnings
                            )

                        if page_text and page_text.strip():
                            extracted_pages.append(page_text)
                            pages_extracted += 1
                        else:
                            # Empty page - might be scanned/image-based
                            warnings.append(
                                f"Page {page_num}: No extractable text (may be image-based)"
                            )
                            logger.warning(
                                f"Empty text on page {page_num} of {self._file_path}"
                            )
                            
                    except Exception as e:
                        warnings.append(f"Page {page_num}: Extraction failed - {str(e)}")
                        logger.error(
                            f"Failed to extract page {page_num} from {self._file_path}: {e}"
                        )

        except pdfplumber.pdfminer.pdfparser.PDFSyntaxError as e:
            return ExtractionResult(
                text=None,
                status=ExtractionStatus.FAILED,
                file_path=str(self._file_path),
                file_type=FileType.PDF,
                error_message=f"Corrupted or invalid PDF: {str(e)}"
            )
        except Exception as e:
            return ExtractionResult(
                text=None,
                status=ExtractionStatus.FAILED,
                file_path=str(self._file_path),
                file_type=FileType.PDF,
                error_message=f"Unexpected error during extraction: {str(e)}"
            )

        # Determine final status
        if pages_extracted == 0:
            status = ExtractionStatus.EMPTY
            combined_text = None
        elif pages_extracted < page_count:
            status = ExtractionStatus.PARTIAL
            combined_text = self.PAGE_SEPARATOR.join(extracted_pages)
        else:
            status = ExtractionStatus.SUCCESS
            combined_text = self.PAGE_SEPARATOR.join(extracted_pages)

        return ExtractionResult(
            text=combined_text,
            status=status,
            file_path=str(self._file_path),
            file_type=FileType.PDF,
            page_count=page_count,
            pages_extracted=pages_extracted,
            warnings=warnings
        )


class DocxTextExtractor(BaseTextExtractor):
    """
    Text extractor for DOCX files using python-docx.
    
    Extracts text from paragraphs and tables while preserving
    basic structure with paragraph separators.
    
    Attributes:
        file_path: Path to the DOCX file.
    """

    PARAGRAPH_SEPARATOR = "\n"

    def __init__(self, file_path: str):
        """
        Initialize the DOCX extractor.
        
        Args:
            file_path: Path to the DOCX file to extract.
        """
        self._file_path = Path(file_path)

    def can_extract(self) -> bool:
        """Check if the DOCX file exists and is readable."""
        return self._file_path.exists() and self._file_path.is_file()

    def extract(self) -> ExtractionResult:
        """
        Extract text from all paragraphs and tables in the DOCX.
        
        Returns:
            ExtractionResult with extracted text and metadata.
            
        Note:
            - Preserves paragraph structure
            - Includes table content
            - Headers and footers are included if present
        """
        if not self.can_extract():
            return ExtractionResult(
                text=None,
                status=ExtractionStatus.FAILED,
                file_path=str(self._file_path),
                file_type=FileType.DOCX,
                error_message=f"File not found or not readable: {self._file_path}"
            )

        warnings: list[str] = []
        extracted_parts: list[str] = []

        try:
            doc = DocxDocument(self._file_path)

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text and para.text.strip():
                    extracted_parts.append(para.text)

            # Extract tables
            for table_idx, table in enumerate(doc.tables, start=1):
                try:
                    table_text = self._extract_table(table)
                    if table_text:
                        extracted_parts.append(table_text)
                except Exception as e:
                    warnings.append(f"Table {table_idx}: Extraction failed - {str(e)}")

        except PackageNotFoundError:
            return ExtractionResult(
                text=None,
                status=ExtractionStatus.FAILED,
                file_path=str(self._file_path),
                file_type=FileType.DOCX,
                error_message="File is not a valid DOCX or is corrupted"
            )
        except Exception as e:
            return ExtractionResult(
                text=None,
                status=ExtractionStatus.FAILED,
                file_path=str(self._file_path),
                file_type=FileType.DOCX,
                error_message=f"Unexpected error during extraction: {str(e)}"
            )

        if not extracted_parts:
            return ExtractionResult(
                text=None,
                status=ExtractionStatus.EMPTY,
                file_path=str(self._file_path),
                file_type=FileType.DOCX,
                page_count=1,
                pages_extracted=0,
                error_message="No extractable text found in document"
            )

        combined_text = self.PARAGRAPH_SEPARATOR.join(extracted_parts)

        return ExtractionResult(
            text=combined_text,
            status=ExtractionStatus.SUCCESS if not warnings else ExtractionStatus.PARTIAL,
            file_path=str(self._file_path),
            file_type=FileType.DOCX,
            page_count=1,  # DOCX doesn't have a direct page count concept
            pages_extracted=1,
            warnings=warnings
        )

    def _extract_table(self, table) -> str:
        """
        Extract text from a table, preserving row/cell structure.
        
        Args:
            table: python-docx Table object.
            
        Returns:
            Formatted table text with tab-separated cells.
        """
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):  # Only include rows with content
                rows.append("\t".join(cells))
        return "\n".join(rows)


class TextExtractorFactory:
    """
    Factory for creating appropriate text extractors based on file type.
    
    This ensures the extraction layer is extensible - new formats can be
    added by implementing BaseTextExtractor and registering here.
    """

    SUPPORTED_EXTENSIONS = {
        ".pdf": FileType.PDF,
        ".docx": FileType.DOCX,
    }

    @classmethod
    def get_extractor(cls, file_path: str) -> BaseTextExtractor:
        """
        Get the appropriate extractor for the given file.
        
        Args:
            file_path: Path to the file to extract.
            
        Returns:
            Appropriate TextExtractor instance.
            
        Raises:
            UnsupportedFileTypeError: If the file type is not supported.
        """
        path = Path(file_path)
        extension = path.suffix.lower()

        if extension not in cls.SUPPORTED_EXTENSIONS:
            supported = ", ".join(cls.SUPPORTED_EXTENSIONS.keys())
            raise UnsupportedFileTypeError(
                f"Unsupported file type: {extension}. Supported: {supported}"
            )

        file_type = cls.SUPPORTED_EXTENSIONS[extension]

        if file_type == FileType.PDF:
            return PdfTextExtractor(file_path)
        elif file_type == FileType.DOCX:
            return DocxTextExtractor(file_path)
        else:
            raise UnsupportedFileTypeError(f"No extractor for type: {file_type}")

    @classmethod
    def is_supported(cls, file_path: str) -> bool:
        """Check if a file type is supported."""
        extension = Path(file_path).suffix.lower()
        return extension in cls.SUPPORTED_EXTENSIONS


def extract_text(file_path: str) -> ExtractionResult:
    """
    Convenience function to extract text from a file.
    
    This is the main entry point for the Text Extraction Layer.
    
    Args:
        file_path: Path to the PDF or DOCX file.
        
    Returns:
        ExtractionResult containing extracted text and metadata.
        
    Example:
        >>> result = extract_text("/path/to/resume.pdf")
        >>> if result.is_success:
        ...     print(result.text)
        ... else:
        ...     print(f"Failed: {result.error_message}")
    """
    try:
        extractor = TextExtractorFactory.get_extractor(file_path)
        return extractor.extract()
    except UnsupportedFileTypeError as e:
        return ExtractionResult(
            text=None,
            status=ExtractionStatus.FAILED,
            file_path=file_path,
            file_type=FileType.PDF,  # Default, actual type unknown
            error_message=str(e)
        )

